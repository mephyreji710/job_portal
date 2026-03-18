"""
Resume Parser — pure Python, no ML libraries.

Pipeline:
    1. Extract raw text  (PDF via pdfplumber, DOCX via python-docx)
    2. Detect sections   (regex header matching)
    3. Extract contact   (email, phone, LinkedIn, GitHub, website)
    4. Match skills      (keyword lookup against built-in lists)
    5. Extract keywords  (frequency after stopword removal)
    6. Estimate experience  (year-range regex)
    7. Detect education level
"""

import io
import re
from datetime import datetime
from collections import Counter

# ---------------------------------------------------------------------------
# Skill lists
# ---------------------------------------------------------------------------

TECH_SKILLS = [
    # Programming languages
    'python', 'javascript', 'typescript', 'java', 'c++', 'c#', 'c', 'go',
    'golang', 'rust', 'kotlin', 'swift', 'ruby', 'php', 'scala', 'r',
    'matlab', 'perl', 'bash', 'shell', 'powershell', 'dart',
    # Web front-end
    'html', 'css', 'html5', 'css3', 'sass', 'less', 'bootstrap',
    'tailwind', 'tailwindcss', 'jquery',
    # Frameworks & libraries
    'django', 'flask', 'fastapi', 'react', 'reactjs', 'angular', 'angularjs',
    'vue', 'vuejs', 'nodejs', 'node.js', 'express', 'expressjs', 'next.js',
    'nuxt.js', 'gatsby', 'spring', 'spring boot', 'laravel', 'rails',
    'ruby on rails', 'asp.net', '.net', 'flutter', 'react native',
    'xamarin', 'ionic', 'electron',
    # Databases
    'sql', 'mysql', 'postgresql', 'postgres', 'sqlite', 'mongodb', 'redis',
    'elasticsearch', 'cassandra', 'dynamodb', 'oracle', 'mssql', 'mariadb',
    'firebase', 'supabase', 'neo4j',
    # Cloud & DevOps
    'aws', 'amazon web services', 'azure', 'gcp', 'google cloud',
    'docker', 'kubernetes', 'k8s', 'jenkins', 'gitlab ci', 'github actions',
    'terraform', 'ansible', 'chef', 'puppet', 'nginx', 'apache', 'linux',
    'unix', 'heroku', 'vercel', 'netlify',
    # ML / AI / Data
    'machine learning', 'deep learning', 'artificial intelligence',
    'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'sklearn',
    'pandas', 'numpy', 'matplotlib', 'seaborn', 'plotly', 'opencv',
    'natural language processing', 'nlp', 'computer vision', 'huggingface',
    # Tools
    'git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'trello',
    'figma', 'photoshop', 'illustrator', 'postman', 'swagger', 'graphql',
    # Concepts / methodologies
    'rest api', 'microservices', 'agile', 'scrum', 'kanban', 'devops',
    'ci/cd', 'tdd', 'bdd', 'oop', 'design patterns', 'mvc', 'mvvm',
    'data structures', 'algorithms', 'cloud computing', 'serverless',
    'blockchain', 'cybersecurity', 'penetration testing',
]

SOFT_SKILLS = [
    'leadership', 'communication', 'teamwork', 'collaboration',
    'problem solving', 'critical thinking', 'time management',
    'project management', 'analytical', 'creativity', 'adaptability',
    'presentation', 'negotiation', 'mentoring', 'coaching', 'multitasking',
    'attention to detail', 'decision making', 'conflict resolution',
    'emotional intelligence', 'interpersonal skills',
]

# Build category lookup (longest first for greedy matching)
_ALL_SKILLS = sorted(
    [(s, 'tech') for s in TECH_SKILLS] + [(s, 'soft') for s in SOFT_SKILLS],
    key=lambda x: len(x[0]), reverse=True,
)

# ---------------------------------------------------------------------------
# Section header patterns
# ---------------------------------------------------------------------------

SECTION_PATTERNS = {
    'summary':        r'\b(summary|objective|profile|about\s+me|professional\s+summary|career\s+objective|executive\s+summary)\b',
    'experience':     r'\b(experience|work\s+experience|employment|professional\s+experience|career\s+history|work\s+history)\b',
    'education':      r'\b(education|academic|qualifications|educational\s+background)\b',
    'skills':         r'\b(skills|technical\s+skills|core\s+competencies|competencies|technologies|tools|expertise)\b',
    'certifications': r'\b(certifications?|certificates?|licenses?|credentials?|courses?|training)\b',
    'projects':       r'\b(projects?|personal\s+projects?|key\s+projects?|portfolio|open\s+source)\b',
}

# ---------------------------------------------------------------------------
# Degree keywords
# ---------------------------------------------------------------------------

DEGREE_MAP = {
    r'\bph\.?d\b|\bdoctorate?\b|\bdoctoral\b': 'phd',
    r"\bmaster'?s?\b|\bm\.sc\b|\bmba\b|\bm\.tech\b|\bmca\b|\bm\.e\b": 'master',
    r"\bbachelor'?s?\b|\bb\.sc\b|\bb\.tech\b|\bb\.e\b|\bbca\b|\bb\.a\b|\bbs\b|\bba\b": 'bachelor',
    r'\bassociate\b': 'associate',
    r'\bdiploma\b|\bcertificate\b': 'diploma',
    r'\bhigh\s+school\b|\bsecondary\b|\bhsc\b|\bssc\b': 'high_school',
}
DEGREE_PRIORITY = ['phd', 'master', 'bachelor', 'associate', 'diploma', 'high_school']

# ---------------------------------------------------------------------------
# Stopwords
# ---------------------------------------------------------------------------

STOPWORDS = {
    'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
    'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through',
    'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had',
    'do', 'does', 'did', 'will', 'would', 'shall', 'should', 'may', 'might',
    'must', 'can', 'could', 'i', 'me', 'my', 'we', 'our', 'you', 'your',
    'he', 'she', 'it', 'they', 'them', 'their', 'this', 'that', 'these',
    'those', 'which', 'who', 'what', 'when', 'where', 'how', 'why', 'all',
    'both', 'each', 'more', 'other', 'some', 'no', 'not', 'only', 'same',
    'so', 'than', 'too', 'very', 'just', 'as', 'if', 'also', 'well', 'new',
    'use', 'used', 'using', 'year', 'years', 'month', 'months', 'day',
    'company', 'team', 'project', 'developed', 'development', 'responsible',
    'experience', 'strong', 'ability', 'knowledge', 'etc', 'including',
    'such', 'over', 'two', 'three', 'one', 'make', 'made', 'good', 'high',
    'help', 'helped', 'worked', 'work', 'working', 'role', 'tasks', 'task',
}


# ---------------------------------------------------------------------------
# ResumeParser
# ---------------------------------------------------------------------------

class ResumeParser:
    """
    Usage:
        parser = ResumeParser(file_bytes, 'pdf')  # or 'docx'
        result = parser.parse()
    """

    def __init__(self, file_content: bytes, file_type: str):
        self.file_content = file_content
        self.file_type = file_type.lower().strip('.')
        self.raw_text = ''

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def parse(self) -> dict:
        self.raw_text = self._extract_text()
        if not self.raw_text.strip():
            raise ValueError("No readable text found in the file. "
                             "The file may be scanned/image-based.")

        sections  = self._detect_sections()
        contact   = self._extract_contact()
        skills    = self._match_skills()
        keywords  = self._extract_keywords()
        exp_years = self._estimate_experience()
        edu_level = self._detect_education_level()

        return {
            'raw_text': self.raw_text,
            'sections': sections,
            'contact': contact,
            'skills': skills,
            'keywords': keywords,
            'estimated_experience_years': exp_years,
            'education_level': edu_level,
        }

    # ------------------------------------------------------------------
    # Step 1 — Text Extraction
    # ------------------------------------------------------------------

    def _extract_text(self) -> str:
        if self.file_type == 'pdf':
            return self._from_pdf()
        if self.file_type in ('docx', 'doc'):
            return self._from_docx()
        raise ValueError(f"Unsupported file type: {self.file_type!r}. "
                         "Please upload a PDF or DOCX file.")

    def _from_pdf(self) -> str:
        import pdfplumber
        parts = []
        with pdfplumber.open(io.BytesIO(self.file_content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if text:
                    parts.append(text)
        return '\n'.join(parts)

    def _from_docx(self) -> str:
        from docx import Document
        doc = Document(io.BytesIO(self.file_content))
        lines = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                lines.append(t)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        lines.append(t)
        return '\n'.join(lines)

    # ------------------------------------------------------------------
    # Step 2 — Section Detection
    # ------------------------------------------------------------------

    def _detect_sections(self) -> dict:
        lines = self.raw_text.split('\n')
        sections = {}
        current = 'header'
        buf = []

        for line in lines:
            stripped = line.strip()
            matched = self._match_section_header(stripped)
            if matched and len(stripped) < 60:   # headers are short lines
                if buf:
                    sections[current] = '\n'.join(buf).strip()
                current = matched
                buf = []
            else:
                buf.append(line)

        if buf:
            sections[current] = '\n'.join(buf).strip()

        return sections

    @staticmethod
    def _match_section_header(line: str) -> str:
        lower = line.lower()
        for section, pattern in SECTION_PATTERNS.items():
            if re.search(pattern, lower):
                return section
        return ''

    # ------------------------------------------------------------------
    # Step 3 — Contact Extraction
    # ------------------------------------------------------------------

    def _extract_contact(self) -> dict:
        text = self.raw_text
        return {
            'name':     self._extract_name(),
            'email':    self._find_email(text),
            'phone':    self._find_phone(text),
            'linkedin': self._find_url(text, r'linkedin\.com/in/[\w\-]+'),
            'github':   self._find_url(text, r'github\.com/[\w\-]+'),
            'website':  self._find_website(text),
        }

    def _extract_name(self) -> str:
        """Heuristic: first non-empty line that looks like a proper name."""
        lines = [l.strip() for l in self.raw_text.split('\n') if l.strip()]
        for line in lines[:6]:
            # 2-4 words, each capitalised, no digits
            if re.match(r'^[A-Z][a-zA-Z]+(?:\s[A-Z][a-zA-Z]+){1,3}$', line):
                return line
        return ''

    @staticmethod
    def _find_email(text: str) -> str:
        m = re.search(r'[\w.%+\-]+@[\w.\-]+\.[a-zA-Z]{2,}', text)
        return m.group(0).lower() if m else ''

    @staticmethod
    def _find_phone(text: str) -> str:
        m = re.search(
            r'(\+?\d{1,3}[\s\-.]?)?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}',
            text)
        return m.group(0).strip() if m else ''

    @staticmethod
    def _find_url(text: str, pattern: str) -> str:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            url = m.group(0)
            return url if url.startswith('http') else 'https://' + url
        return ''

    @staticmethod
    def _find_website(text: str) -> str:
        """Find a personal website URL (not LinkedIn/GitHub)."""
        m = re.search(
            r'https?://(?!.*linkedin)(?!.*github)[\w\-]+\.[\w\-./]+',
            text, re.IGNORECASE)
        return m.group(0) if m else ''

    # ------------------------------------------------------------------
    # Step 4 — Skills Matching
    # ------------------------------------------------------------------

    def _match_skills(self) -> list:
        text_lower = self.raw_text.lower()
        found = []
        seen = set()
        for skill, category in _ALL_SKILLS:
            if skill in seen:
                continue
            pattern = r'(?<![a-z])' + re.escape(skill) + r'(?![a-z])'
            if re.search(pattern, text_lower):
                found.append({'name': skill, 'category': category})
                seen.add(skill)
        return found

    # ------------------------------------------------------------------
    # Step 5 — Keyword Extraction
    # ------------------------------------------------------------------

    def _extract_keywords(self, top_n: int = 40) -> list:
        # strip non-alpha
        clean = re.sub(r'[^a-zA-Z\s]', ' ', self.raw_text.lower())
        words = [w for w in clean.split()
                 if w not in STOPWORDS and len(w) > 2 and not w.isdigit()]
        counter = Counter(words)
        return [{'word': w, 'count': c}
                for w, c in counter.most_common(top_n)]

    # ------------------------------------------------------------------
    # Step 6 — Experience Estimation
    # ------------------------------------------------------------------

    def _estimate_experience(self) -> float | None:
        pattern = (
            r'\b(20\d{2}|19\d{2})\b'
            r'\s*[-–—to]+\s*'
            r'\b(20\d{2}|19\d{2}|present|current|now|till\s+date)\b'
        )
        current_year = datetime.now().year
        total_months = 0
        seen_ranges = set()

        for m in re.finditer(pattern, self.raw_text, re.IGNORECASE):
            start = int(m.group(1))
            end_s = m.group(2).lower()
            end   = current_year if re.search(r'present|current|now|date', end_s) else int(end_s)

            key = (start, end)
            if key in seen_ranges:
                continue
            seen_ranges.add(key)

            if start <= end <= current_year + 1 and (end - start) <= 40:
                total_months += (end - start) * 12

        return round(total_months / 12, 1) if total_months else None

    # ------------------------------------------------------------------
    # Step 7 — Education Level Detection
    # ------------------------------------------------------------------

    def _detect_education_level(self) -> str:
        text_lower = self.raw_text.lower()
        for level in DEGREE_PRIORITY:
            for pattern, mapped in DEGREE_MAP.items():
                if mapped == level and re.search(pattern, text_lower):
                    return level
        return ''


# ---------------------------------------------------------------------------
# Helper — called from views
# ---------------------------------------------------------------------------

def run_parse_and_save(resume_obj) -> 'ParsedResume':
    """
    Read file, parse, persist to ParsedResume + ExtractedKeyword.
    Returns the ParsedResume instance.
    Raises on error (caller should catch and mark resume as failed).
    """
    from django.utils import timezone
    from .models import ParsedResume, ExtractedKeyword

    # Read bytes from storage
    resume_obj.file.seek(0)
    content = resume_obj.file.read()

    parser = ResumeParser(content, resume_obj.file_type)
    result = parser.parse()

    # Persist raw text + status
    resume_obj.raw_text = result['raw_text']
    resume_obj.status   = 'parsed'
    resume_obj.parsed_at = timezone.now()
    resume_obj.save(update_fields=['raw_text', 'status', 'parsed_at'])

    skills   = result['skills']
    keywords = result['keywords']
    sections = result['sections']
    contact  = result['contact']

    tech_count = sum(1 for s in skills if s['category'] == 'tech')
    soft_count = sum(1 for s in skills if s['category'] == 'soft')

    parsed, _ = ParsedResume.objects.update_or_create(
        resume=resume_obj,
        defaults={
            'extracted_name':    contact.get('name', ''),
            'extracted_email':   contact.get('email', ''),
            'extracted_phone':   contact.get('phone', ''),
            'extracted_linkedin': contact.get('linkedin', ''),
            'extracted_github':   contact.get('github', ''),
            'extracted_website':  contact.get('website', ''),

            'summary_text':        sections.get('summary', ''),
            'skills_text':         sections.get('skills', ''),
            'education_text':      sections.get('education', ''),
            'experience_text':     sections.get('experience', ''),
            'certifications_text': sections.get('certifications', ''),
            'projects_text':       sections.get('projects', ''),

            'extracted_skills':   skills,
            'extracted_keywords': keywords[:20],

            'estimated_experience_years': result.get('estimated_experience_years'),
            'education_level':    result.get('education_level', ''),
            'total_words':        len(result['raw_text'].split()),
            'total_keywords':     len(keywords),
            'tech_skill_count':   tech_count,
            'soft_skill_count':   soft_count,
        },
    )

    # Save individual keywords
    ExtractedKeyword.objects.filter(parsed_resume=parsed).delete()
    skill_names = {s['name'] for s in skills}
    kw_objs = []
    for kw in keywords[:60]:
        w = kw['word']
        is_skill = w in skill_names
        cat = next((s['category'] for s in skills if s['name'] == w), 'other')
        kw_objs.append(ExtractedKeyword(
            parsed_resume=parsed,
            word=w,
            frequency=kw['count'],
            is_skill=is_skill,
            category=cat,
        ))
    ExtractedKeyword.objects.bulk_create(kw_objs, ignore_conflicts=True)

    return parsed
